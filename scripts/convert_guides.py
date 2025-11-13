"""Convert Word drug guides into JSON data for the static website.

Usage:
    python convert_guides.py --source "../Drug Guides" --output "../data"

The script expects each course to exist as a folder inside the source directory.
Every `.docx` inside a course folder is converted to JSON and an HTML fragment,
and `guides.index.json` in the output directory is refreshed. Optional
`tags.txt` files inside course folders provide a list of tags that are attached
to every guide in that course.
"""

from __future__ import annotations

import argparse
import datetime as dt
import json
import re
import sys
from dataclasses import dataclass
from html import escape
from pathlib import Path
from typing import Iterable, List, Sequence

from docx import Document  # type: ignore
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore
from docx.oxml.ns import qn  # type: ignore


# Get the script's directory to resolve relative paths
SCRIPT_DIR = Path(__file__).parent.resolve()
DEFAULT_SOURCE = SCRIPT_DIR.parent / "Drug Guides"
DEFAULT_OUTPUT = SCRIPT_DIR.parent / "data"
INDEX_FILENAME = "guides.index.json"
HTML_SUBDIR = "html"
DATE_FORMAT = "%Y-%m-%dT%H:%M:%S%z"


@dataclass
class CourseContext:
    name: str
    slug: str
    path: Path
    tags: List[str]


@dataclass
class GuideMetadata:
    title: str
    course: CourseContext
    slug: str
    source_path: Path
    json_path: Path
    html_path: Path
    tags: List[str]
    table_count: int = 0


def parse_args(argv: Sequence[str]) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Convert Word drug guides to JSON.")
    parser.add_argument(
        "--source",
        type=Path,
        default=DEFAULT_SOURCE,
        help="Directory containing Word documents (default: ../Drug Guides).",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=DEFAULT_OUTPUT,
        help="Directory to write JSON files (default: ../data).",
    )
    parser.add_argument(
        "--index",
        type=str,
        default=INDEX_FILENAME,
        help=f"Index filename relative to the output directory (default: {INDEX_FILENAME}).",
    )
    return parser.parse_args(argv)


def slugify(text: str) -> str:
    normalized = re.sub(r"[^A-Za-z0-9]+", "-", text).strip("-")
    return normalized.lower()


def parse_metadata(doc_path: Path, output_dir: Path, course: CourseContext) -> GuideMetadata:
    stem = doc_path.stem
    title = stem.replace("-", " ").replace("_", " ").strip() or stem

    document_slug = slugify(stem)
    slug = slugify(f"{course.slug}-{document_slug}") if course.slug else document_slug

    json_path = output_dir / f"{slug}.json"
    html_path = output_dir / HTML_SUBDIR / f"{slug}.html"

    return GuideMetadata(
        title=title,
        course=course,
        slug=slug,
        source_path=doc_path,
        json_path=json_path,
        html_path=html_path,
        tags=list(course.tags),
    )


def extract_text(cell) -> str:
    lines: List[str] = []
    for paragraph in cell.paragraphs:
        text = "".join(run.text for run in paragraph.runs).strip()
        if text:
            lines.append(text)
    return "<br>".join(lines).strip()


def convert_table(table) -> dict:
    headers: List[str] = []
    rows = []

    if table.rows:
        first_row = table.rows[0]
        headers = [extract_text(cell) for cell in first_row.cells]

        data_rows = table.rows[1:]
        for row in data_rows:
            rows.append([extract_text(cell) for cell in row.cells])

    return {
        "headers": headers,
        "rows": rows,
    }


def generate_cell_data(tables: List[dict]) -> dict:
    """Generate cellData structure for all non-empty cells in tables.
    
    Uses 1-based table indices to match HTML data-table-index attributes.
    """
    cell_data = {}
    
    for table_idx, table in enumerate(tables, start=1):  # Start at 1 to match HTML
        # Process headers (row 0)
        if table.get("headers"):
            for col_idx, header in enumerate(table["headers"]):
                if header and header.strip() and header.strip() != "&nbsp;":
                    cell_id = f"table_{table_idx}_row_0_col_{col_idx}"
                    # Normalize content: remove HTML tags, strip whitespace
                    normalized = re.sub(r"<[^>]+>", "", header).strip()
                    if normalized:
                        cell_data[cell_id] = {
                            "content": normalized,
                            "summary": "",
                        }
        
        # Process rows
        if table.get("rows"):
            for row_idx, row in enumerate(table["rows"], start=1):
                for col_idx, cell in enumerate(row):
                    if cell and cell.strip() and cell.strip() != "&nbsp;":
                        cell_id = f"table_{table_idx}_row_{row_idx}_col_{col_idx}"
                        # Normalize content: remove HTML tags, strip whitespace
                        normalized = re.sub(r"<[^>]+>", "", cell).strip()
                        if normalized:
                            cell_data[cell_id] = {
                                "content": normalized,
                                "summary": "",
                            }
    
    return cell_data


def convert_document(meta: GuideMetadata, document: Document) -> dict:
    tables = [convert_table(table) for table in document.tables if table.rows]
    meta.table_count = len(tables)

    # Preserve existing cellData if JSON file already exists
    existing_cell_data = {}
    if meta.json_path.exists():
        try:
            with meta.json_path.open("r", encoding="utf-8") as fh:
                existing_data = json.load(fh)
                existing_cell_data = existing_data.get("cellData", {})
        except (json.JSONDecodeError, IOError):
            pass

    # Generate new cellData for all non-empty cells
    new_cell_data = generate_cell_data(tables)
    
    # Build a content-to-summary map from existing data for migration
    # This helps preserve summaries when cell IDs change (e.g., 0-based to 1-based indices)
    content_to_summary = {}
    for existing_entry in existing_cell_data.values():
        content = existing_entry.get("content", "").strip()
        summary = existing_entry.get("summary", "").strip()
        if content and summary and summary != "no data":
            # Use content as key (normalized)
            normalized_content = re.sub(r"<[^>]+>", "", content).strip()
            if normalized_content:
                content_to_summary[normalized_content] = {
                    "summary": summary,
                    "lastUpdated": existing_entry.get("lastUpdated", ""),
                }
    
    # Merge: preserve existing entries, add new ones
    merged_cell_data = {**new_cell_data}
    for cell_id, new_entry in new_cell_data.items():
        # First try to match by cell_id (exact match)
        if cell_id in existing_cell_data:
            existing_entry = existing_cell_data[cell_id]
            if existing_entry.get("summary"):
                merged_cell_data[cell_id]["summary"] = existing_entry["summary"]
            if "lastUpdated" in existing_entry:
                merged_cell_data[cell_id]["lastUpdated"] = existing_entry["lastUpdated"]
        else:
            # If cell_id doesn't match, try to match by content (for index migration)
            content = new_entry.get("content", "").strip()
            normalized_content = re.sub(r"<[^>]+>", "", content).strip()
            if normalized_content in content_to_summary:
                # Found a match by content - migrate the summary
                matched_data = content_to_summary[normalized_content]
                merged_cell_data[cell_id]["summary"] = matched_data["summary"]
                if matched_data.get("lastUpdated"):
                    merged_cell_data[cell_id]["lastUpdated"] = matched_data["lastUpdated"]

    if not tables:
        return {
            "title": meta.title,
            "course": meta.course.name,
            "courseSlug": meta.course.slug,
            "tags": meta.tags,
            "tables": [],
            "cellData": merged_cell_data,
        }

    return {
        "title": meta.title,
        "course": meta.course.name,
        "courseSlug": meta.course.slug,
        "tags": meta.tags,
        "tables": tables,
        "cellData": merged_cell_data,
    }


def build_html_fragment(meta: GuideMetadata, document: Document) -> str:
    table_parts: List[str] = []

    for index, table in enumerate(document.tables, start=1):
        table_parts.append(render_table_html(table, index))

    if not table_parts:
        table_parts.append(
            '<p class="guide-empty">No tables were found in this guide.</p>'
        )

    inner_html = "\n".join(table_parts)
    return f'<section class="guide-fragment" data-guide="{meta.slug}">\n{inner_html}\n</section>'


def render_table_html(table, index: int) -> str:
    rows_html: List[str] = []
    max_columns = max(len(row.cells) for row in table.rows) if table.rows else 0

    for row_idx, row in enumerate(table.rows):
        cell_tag = "th" if row_idx == 0 else "td"
        rendered_cells = set()
        cell_chunks: List[str] = []

        for col_idx, cell in enumerate(row.cells):
            cell_id = id(cell._tc)
            if cell_id in rendered_cells:
                continue
            rendered_cells.add(cell_id)

            rowspan = get_rowspan(table, row_idx, col_idx, cell)
            if rowspan == 0:
                continue
            colspan = get_colspan(cell)

            cell_styles = collect_cell_styles(cell)
            style_attr = f' style="{cell_styles}"' if cell_styles else ""

            attrs: List[str] = []
            if colspan > 1:
                attrs.append(f'colspan="{colspan}"')
            if rowspan > 1:
                attrs.append(f'rowspan="{rowspan}"')
            if style_attr:
                attrs.append(style_attr.strip())

            cell_html = render_cell_html(cell)
            attr_string = (" " + " ".join(attrs)) if attrs else ""
            cell_chunks.append(f"    <{cell_tag}{attr_string}>{cell_html}</{cell_tag}>")

        if cell_chunks:
            rows_html.append("  <tr>\n" + "\n".join(cell_chunks) + "\n  </tr>")

    table_attrs = [
        f'class="guide-table guide-table-{index}"',
        f'data-table-index="{index}"',
    ]
    if max_columns:
        table_attrs.append(f'data-columns="{max_columns}"')
    table_style = collect_table_styles(table)
    if table_style:
        table_attrs.append(f'style="{table_style}"')

    return "<table {attrs}>\n{rows}\n</table>".format(
        attrs=" ".join(table_attrs),
        rows="\n".join(rows_html),
    )


def render_cell_html(cell) -> str:
    paragraphs_html: List[str] = []
    for paragraph in cell.paragraphs:
        paragraphs_html.append(render_paragraph_html(paragraph))

    merged_text = "\n".join(paragraphs_html)
    return merged_text or "&nbsp;"


def render_paragraph_html(paragraph) -> str:
    runs_html = []
    for run in paragraph.runs:
        run_html = render_run_html(run)
        if run_html:
            runs_html.append(run_html)

    text_html = "".join(runs_html)
    if not text_html:
        text_html = "&nbsp;"

    styles: List[str] = []
    alignment = paragraph.alignment
    if alignment is not None:
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            styles.append("text-align: center;")
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            styles.append("text-align: right;")
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            styles.append("text-align: justify;")

    para_font = getattr(paragraph.style, "font", None) if paragraph.style else None
    if para_font is not None:
        size = getattr(para_font, "size", None)
        if size and hasattr(size, "pt"):
            styles.append(f"font-size: {size.pt:.2f}pt;")

        font_color = getattr(para_font, "color", None)
        if font_color and font_color.rgb:
            styles.append(f"color: #{font_color.rgb};")

    para_format = paragraph.paragraph_format
    if para_format is not None:
        before = getattr(para_format, "space_before", None)
        after = getattr(para_format, "space_after", None)
        line_space = getattr(para_format, "line_spacing", None)
        if before and hasattr(before, "pt"):
            styles.append(f"margin-top: {before.pt:.2f}pt;")
        if after and hasattr(after, "pt"):
            styles.append(f"margin-bottom: {after.pt:.2f}pt;")
        if line_space and isinstance(line_space, (int, float)):
            styles.append(f"line-height: {line_space};")

    class_names: List[str] = []
    if paragraph.style and paragraph.style.name:
        sanitized_name = slugify(paragraph.style.name)
        if sanitized_name:
            class_names.append(f"para-{sanitized_name}")

    if paragraph._p.pPr is not None and paragraph._p.pPr.numPr is not None:
        class_names.append("para-list")

    class_attr = f' class="{" ".join(class_names)}"' if class_names else ""
    style_attr = f' style="{" ".join(styles)}"' if styles else ""

    return f"<p{class_attr}{style_attr}>{text_html}</p>"


def render_run_html(run) -> str:
    text = run.text or ""
    if not text:
        return ""

    text = escape(text)

    open_tags: List[str] = []
    close_tags: List[str] = []

    if run.bold:
        open_tags.append("<strong>")
        close_tags.insert(0, "</strong>")

    if run.italic:
        open_tags.append("<em>")
        close_tags.insert(0, "</em>")

    if run.underline:
        open_tags.append('<span style="text-decoration: underline;">')
        close_tags.insert(0, "</span>")

    style_bits: List[str] = []
    if run.font and run.font.color and run.font.color.rgb:
        style_bits.append(f"color: #{run.font.color.rgb};")

    if run.font and run.font.highlight_color:
        highlight_key = getattr(run.font.highlight_color, "name", str(run.font.highlight_color))
        highlight = COLOR_INDEX_TO_HEX.get(highlight_key)
        if highlight:
            style_bits.append(f"background-color: {highlight};")

    if style_bits:
        open_tags.append(f'<span style="{" ".join(style_bits)}">')
        close_tags.insert(0, "</span>")

    return "".join(open_tags) + text + "".join(close_tags)


COLOR_INDEX_TO_HEX = {
    "YELLOW": "#fff200",
    "TURQUOISE": "#00a8e8",
    "BRIGHT_GREEN": "#66ff00",
    "PINK": "#ff66cc",
    "BLUE": "#4f81bd",
    "RED": "#ff0000",
    "DARK_BLUE": "#17365d",
    "TEAL": "#31859b",
    "GREEN": "#00b050",
    "VIOLET": "#7030a0",
    "DARK_RED": "#c00000",
    "DARK_YELLOW": "#806000",
    "GRAY_50": "#808080",
    "GRAY_25": "#c0c0c0",
    "BLACK": "#000000",
}

BORDER_STYLE_MAP = {
    "nil": "none",
    "none": "none",
    "single": "solid",
    "double": "double",
    "dashed": "dashed",
    "dotted": "dotted",
    "thick": "solid",
    "hairline": "solid",
    "wave": "wavy",
    "dashSmallGap": "dashed",
    "dashDot": "dashed",
    "dashDotDot": "dashed",
    "triple": "double",
}


def twips_to_pt(value: int | None) -> float | None:
    if value is None:
        return None
    return value / 20.0


def eighth_pt_to_pt(value: int | None) -> float | None:
    if value is None:
        return None
    return value / 8.0


def build_style_string(styles: Iterable[str]) -> str:
    filtered = [style.strip().rstrip(";") for style in styles if style]
    if not filtered:
        return ""
    return "; ".join(filtered) + ";"


def get_colspan(cell) -> int:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if tc_pr is None or tc_pr.gridSpan is None or tc_pr.gridSpan.val is None:
        return 1
    try:
        return int(tc_pr.gridSpan.val)
    except (TypeError, ValueError):
        return 1


def get_rowspan(table, row_idx: int, col_idx: int, cell) -> int:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if tc_pr is None or tc_pr.vMerge is None:
        return 1

    v_val = tc_pr.vMerge.val
    if v_val is not None and v_val != "restart":
        return 0

    span = 1
    next_row_idx = row_idx + 1
    while next_row_idx < len(table.rows):
        try:
            next_cell = table.rows[next_row_idx].cells[col_idx]
        except IndexError:
            break
        next_tc_pr = getattr(next_cell._tc, "tcPr", None)
        if next_tc_pr is None or next_tc_pr.vMerge is None:
            break
        next_val = next_tc_pr.vMerge.val
        if next_val not in (None, "continue"):
            break
        span += 1
        next_row_idx += 1
    return span


def get_cell_borders(cell) -> dict[str, str]:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if tc_pr is None:
        return {}

    try:
        tc_borders = tc_pr.tcBorders
    except AttributeError:
        tc_borders = None

    if tc_borders is None:
        borders_el = tc_pr.xpath("./w:tcBorders")
        tc_borders = borders_el[0] if borders_el else None

    if tc_borders is None:
        return {}

    def get_border_element(border_container, side_name):
        if hasattr(border_container, side_name):
            try:
                return getattr(border_container, side_name)
            except AttributeError:
                pass
        return border_container.find(qn(f"w:{side_name}")) if hasattr(border_container, "find") else None

    borders = {}
    border_map = {
        "top": "border-top",
        "bottom": "border-bottom",
        "left": "border-left",
        "right": "border-right",
    }
    for side, css_name in border_map.items():
        border = get_border_element(tc_borders, side)
        if border is None:
            continue
        style_key = getattr(border, "val", None) if hasattr(border, "val") else border.get(qn("w:val"))
        if style_key in ("nil", "none"):
            borders[css_name] = "none"
            continue
        border_style = BORDER_STYLE_MAP.get(style_key, "solid")
        size_raw = getattr(border, "sz", None) if hasattr(border, "sz") else border.get(qn("w:sz"))
        try:
            size_value = int(size_raw) if size_raw is not None else None
        except (TypeError, ValueError):
            size_value = None
        size_pt = eighth_pt_to_pt(size_value)
        width_part = f"{size_pt:.2f}pt" if size_pt else "1pt"
        color = getattr(border, "color", None) if hasattr(border, "color") else border.get(qn("w:color"))
        color_part = f"#{color}" if color and color != "auto" else "#13294b"
        borders[css_name] = f"{width_part} {border_style} {color_part}"
    return borders


def get_cell_padding(cell) -> dict[str, str]:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if tc_pr is None:
        return {}

    try:
        tc_mar = tc_pr.tcMar
    except AttributeError:
        tc_mar = None

    if tc_mar is None and tc_pr is not None:
        mar_elems = tc_pr.xpath("./w:tcMar")
        tc_mar = mar_elems[0] if mar_elems else None

    if tc_mar is None:
        return {}

    padding = {}
    margin_map = {
        "top": "padding-top",
        "bottom": "padding-bottom",
        "left": "padding-left",
        "right": "padding-right",
    }
    for side, css_name in margin_map.items():
        if hasattr(tc_mar, side):
            margin = getattr(tc_mar, side, None)
        else:
            margin = tc_mar.find(qn(f"w:{side}")) if hasattr(tc_mar, "find") else None
        if margin is None:
            continue
        value = getattr(margin, "w", None) if hasattr(margin, "w") else margin.get(qn("w:w"))
        if value is None:
            continue
        try:
            value_int = int(value)
        except (TypeError, ValueError):
            continue
        pt_value = twips_to_pt(value_int)
        if pt_value is not None:
            padding[css_name] = f"{pt_value:.2f}pt"
    return padding


def collect_cell_styles(cell) -> str:
    styles: List[str] = []
    background = get_cell_background(cell)
    if background:
        styles.append(f"background-color: {background}")
    vertical = get_cell_vertical_alignment(cell)
    if vertical:
        styles.append(f"vertical-align: {vertical}")

    for css_name, value in {**get_cell_borders(cell), **get_cell_padding(cell)}.items():
        styles.append(f"{css_name}: {value}")

    return build_style_string(styles)


def collect_table_styles(table) -> str:
    styles: List[str] = ["border-collapse: collapse"]
    tbl_pr = getattr(table._tbl, "tblPr", None)
    tbl_borders = None
    if tbl_pr is not None:
        try:
            tbl_borders = tbl_pr.tblBorders
        except AttributeError:
            borders_el = tbl_pr.xpath("./w:tblBorders")
            tbl_borders = borders_el[0] if borders_el else None

    if tbl_borders is not None:
        border_values = []
        for side in ("top", "bottom", "left", "right", "insideH", "insideV"):
            if hasattr(tbl_borders, side):
                border = getattr(tbl_borders, side, None)
            else:
                border = tbl_borders.find(qn(f"w:{side}")) if hasattr(tbl_borders, "find") else None
            if border is None:
                continue
            style_key = getattr(border, "val", None) if hasattr(border, "val") else border.get(qn("w:val"))
            if style_key in ("nil", "none"):
                continue
            css_style = BORDER_STYLE_MAP.get(style_key, "solid")
            raw_size = getattr(border, "sz", None) if hasattr(border, "sz") else border.get(qn("w:sz"))
            size_pt = eighth_pt_to_pt(int(raw_size)) if raw_size is not None else None
            width_part = f"{size_pt:.2f}pt" if size_pt else "1pt"
            color = getattr(border, "color", None) if hasattr(border, "color") else border.get(qn("w:color"))
            color_part = f"#{color}" if color and color != "auto" else "#13294b"
            border_values.append(f"{width_part} {css_style} {color_part}")
        if border_values:
            styles.append(f"border: {border_values[0]}")

    tbl_cell_mar = None
    if tbl_pr is not None:
        try:
            tbl_cell_mar = tbl_pr.tblCellMar
        except AttributeError:
            mar_el = tbl_pr.xpath("./w:tblCellMar")
            tbl_cell_mar = mar_el[0] if mar_el else None

    if tbl_cell_mar is not None:
        spacing_values = []
        for side in ("top", "bottom", "left", "right"):
            if hasattr(tbl_cell_mar, side):
                margin = getattr(tbl_cell_mar, side, None)
            else:
                margin = tbl_cell_mar.find(qn(f"w:{side}")) if hasattr(tbl_cell_mar, "find") else None
            if margin is None:
                continue
            value = margin.w if hasattr(margin, "w") else margin.get(qn("w:w"), None) if hasattr(margin, "get") else None
            if value is None:
                continue
            if hasattr(value, "val"):
                value = value.val
            try:
                value_int = int(value)
            except (TypeError, ValueError):
                continue
            pt_value = twips_to_pt(value_int)
            if pt_value is not None:
                spacing_values.append(pt_value)
        if spacing_values:
            styles.append(f"border-spacing: {max(spacing_values):.2f}pt")

    return build_style_string(styles)

BORDER_STYLE_MAP = {
    "nil": "none",
    "none": "none",
    "single": "solid",
    "double": "double",
    "dashed": "dashed",
    "dotted": "dotted",
    "thick": "solid",
    "hairline": "solid",
    "wave": "wavy",
    "dashSmallGap": "dashed",
    "dashDot": "dashed",
    "dashDotDot": "dashed",
    "triple": "double",
}


def get_cell_background(cell) -> str | None:
    tc_pr = getattr(cell._tc, "tcPr", None)
    if tc_pr is None:
        return None
    shd_elems = tc_pr.xpath("./w:shd")
    if shd_elems:
        fill = shd_elems[0].get(qn("w:fill"))
        if fill and fill != "auto":
            return f"#{fill}"
    return None


def get_cell_vertical_alignment(cell) -> str | None:
    v_align = cell.vertical_alignment
    if v_align is None:
        return None
    try:
        if v_align == WD_CELL_VERTICAL_ALIGNMENT.TOP:
            return "top"
        if v_align == WD_CELL_VERTICAL_ALIGNMENT.CENTER:
            return "middle"
        if v_align == WD_CELL_VERTICAL_ALIGNMENT.BOTTOM:
            return "bottom"
    except Exception:
        pass
    str_value = str(v_align).lower()
    if "center" in str_value:
        return "middle"
    if "bottom" in str_value:
        return "bottom"
    if "top" in str_value:
        return "top"
    return None


def write_json(path: Path, payload: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        json.dump(payload, fh, ensure_ascii=False, indent=2)
        fh.write("\n")


def write_html(path: Path, markup: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8") as fh:
        fh.write(markup)
        if not markup.endswith("\n"):
            fh.write("\n")


def update_index(
    output_dir: Path,
    index_filename: str,
    entries: Iterable[GuideMetadata],
) -> None:
    index_payload = {
        "generated": dt.datetime.now(dt.timezone.utc).strftime(DATE_FORMAT),
        "guides": [
            {
                "title": meta.title,
                "course": meta.course.name,
                "courseSlug": meta.course.slug,
                "tags": meta.tags,
                "slug": meta.slug,
                "dataFile": meta.json_path.name,
                "fragment": meta.html_path.relative_to(output_dir).as_posix(),
                "sourceFile": str(meta.source_path),
                "tables": meta.table_count,
            }
            for meta in sorted(entries, key=lambda m: m.title.lower())
        ],
    }

    write_json(output_dir / index_filename, index_payload)


def read_course_tags(course_dir: Path) -> List[str]:
    tags_path = course_dir / "tags.txt"
    if not tags_path.exists():
        return []

    tags: List[str] = []
    with tags_path.open("r", encoding="utf-8") as fh:
        for line in fh:
            tag = line.strip()
            if tag:
                tags.append(tag)

    # Preserve first occurrence order but sort case-insensitively for stability.
    unique_in_order: List[str] = []
    seen = set()
    for tag in tags:
        normalized = tag.lower()
        if normalized in seen:
            continue
        seen.add(normalized)
        unique_in_order.append(tag)

    return sorted(unique_in_order, key=lambda t: t.lower())


def discover_courses(source_dir: Path) -> List[CourseContext]:
    course_dirs = [path for path in source_dir.iterdir() if path.is_dir()]
    courses: List[CourseContext] = []

    for directory in sorted(course_dirs, key=lambda p: p.name.lower()):
        course_name = directory.name.replace("_", " ").strip() or directory.name
        course_slug = slugify(course_name)
        tags = read_course_tags(directory)
        courses.append(
            CourseContext(
                name=course_name,
                slug=course_slug,
                path=directory,
                tags=tags,
            )
        )

    return courses


def main(argv: Sequence[str] | None = None) -> int:
    args = parse_args(argv or sys.argv[1:])

    source_dir: Path = args.source.resolve()
    output_dir: Path = args.output.resolve()

    if not source_dir.exists():
        print(f"Source directory not found: {source_dir}", file=sys.stderr)
        return 1

    courses = discover_courses(source_dir)
    if not courses:
        print(f"No course folders found in {source_dir}", file=sys.stderr)
        return 1

    metadata_entries: List[GuideMetadata] = []

    for course in courses:
        documents = sorted(course.path.glob("*.docx"))
        if not documents:
            print(f"Skipping {course.name}: no .docx files found.", file=sys.stderr)
            continue

        for doc_path in documents:
            document = Document(doc_path)
            meta = parse_metadata(doc_path, output_dir, course)
            payload = convert_document(meta, document)
            write_json(meta.json_path, payload)

            html_fragment = build_html_fragment(meta, document)
            write_html(meta.html_path, html_fragment)

            metadata_entries.append(meta)
            print(
                f"Converted {doc_path.relative_to(course.path)} -> {meta.json_path.name} "
                f"and {meta.html_path.relative_to(output_dir)} ({meta.table_count} tables) "
                f"[course={course.name}]"
            )

    if not metadata_entries:
        print("No guides converted. Ensure each course folder contains .docx files.", file=sys.stderr)
        return 1

    update_index(output_dir, args.index, metadata_entries)
    print(f"Updated index at {output_dir / args.index}")
    return 0


if __name__ == "__main__":
    exit_code = main()
    # If running directly (double-clicked), pause on error so user can see the message
    if exit_code != 0:
        input("\nPress Enter to exit...")
    raise SystemExit(exit_code)

